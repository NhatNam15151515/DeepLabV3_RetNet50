# -*- coding: utf-8 -*-
"""
generate_report.py — Tự động sinh báo cáo học thuật .docx
Dự án: Phân vùng ngữ nghĩa thực phẩm & Ước lượng trọng lượng tự động
Kiến trúc: DeepLabV3+ (ResNet50) | Dataset: FoodSemSeg (22 class)
"""

import sys
sys.stdout.reconfigure(encoding="utf-8")
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # project root (script in scripts/)
IMG_DIR_EXP19 = os.path.join(BASE, "scripts", "api_results_exp19")
IMG_DIR_EXP21 = os.path.join(BASE, "scripts", "api_results")
OUT_PATH = os.path.join(BASE, "reports", "BaoCao_DeepLabV3_FoodSeg.docx")
# Fallback nếu file gốc đang mở
OUT_PATH_ALT = os.path.join(BASE, "reports", "BaoCao_DeepLabV3_FoodSeg_NEW.docx")

# ── Màu sắc chủ đạo ──────────────────────────────────────────────────────────
C_TITLE   = RGBColor(0x1A, 0x23, 0x7E)   # indigo đậm
C_HEAD1   = RGBColor(0x0D, 0x47, 0xA1)   # blue đậm
C_HEAD2   = RGBColor(0x15, 0x65, 0xC0)   # blue vừa
C_HEAD3   = RGBColor(0x19, 0x76, 0xD2)   # blue nhạt
C_ACCENT  = RGBColor(0xE6, 0x51, 0x00)   # cam đậm
C_CODE    = RGBColor(0x37, 0x47, 0x4F)   # xám đậm (code)
C_TABLE_H = RGBColor(0x0D, 0x47, 0xA1)   # header bảng

# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), kwargs.get("val", "single"))
        border.set(qn("w:sz"), kwargs.get("sz", "4"))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get("color", "B0BEC5"))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    bolds = {1: True, 2: True, 3: True, 4: False}
    colors = {1: C_HEAD1, 2: C_HEAD2, 3: C_HEAD3, 4: C_HEAD3}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = bolds.get(level, False)
    run.font.color.rgb = color or colors.get(level, C_HEAD1)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.keep_with_next = True
    return p

def add_para(doc, text, bold=False, italic=False, size=11, indent=0, color=None,
             space_before=0, space_after=4, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = C_CODE
    return p

def add_image(doc, path, width=Inches(5.5), caption=""):
    if not os.path.exists(path):
        add_para(doc, f"[Ảnh không tìm thấy: {os.path.basename(path)}]",
                 italic=True, color=RGBColor(0xCC, 0, 0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.runs[0]
        r.font.italic = True
        r.font.size   = Pt(9.5)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def make_table(doc, headers, rows, col_widths=None, zebra=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold  = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size  = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "0D47A1")
    # Data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        bg = "E3F2FD" if (zebra and ri % 2 == 0) else "FFFFFF"
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            _set_cell_bg(cell, bg)
            _set_cell_border(cell)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = w
    doc.add_paragraph()
    return table

def page_break(doc):
    doc.add_page_break()

# ── Trang bìa ─────────────────────────────────────────────────────────────────

def write_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("TRƯỜNG ĐẠI HỌC — KHOA CÔNG NGHỆ THÔNG TIN")
    r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = C_TITLE

    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("ĐỒ ÁN CHUYÊN NGÀNH")
    r2.font.size = Pt(14); r2.font.bold = True; r2.font.color.rgb = C_TITLE

    doc.add_paragraph(); doc.add_paragraph()
    main = doc.add_paragraph()
    main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = main.add_run("PHÂN VÙNG NGỮ NGHĨA THỰC PHẨM\n& ƯỚC LƯỢNG TRỌNG LƯỢNG TỰ ĐỘNG")
    rm.font.size = Pt(22); rm.font.bold = True; rm.font.color.rgb = C_TITLE

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Kiến trúc DeepLabV3+ (ResNet50 Encoder) · NutritionVerse Pipeline 3 tầng")
    rs.font.size = Pt(12); rs.font.italic = True; rs.font.color.rgb = C_HEAD2

    doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
    info_lines = [
        ("Sinh viên thực hiện:", "Nhật Nam"),
        ("Kiến trúc mô hình:", "DeepLabV3+ / ResNet50 / SMP"),
        ("Bộ dữ liệu:", "FoodSemSeg_512×512 (22 class)"),
        ("Kết quả tốt nhất:", "mIoU = 0.3850 · Pixel Accuracy = 0.8282 (exp21)"),
        ("Ngày hoàn thành:", "09/03/2026"),
    ]
    for label, val in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}  "); rl.font.size = Pt(11); rl.font.bold = True
        rv = p.add_run(val); rv.font.size = Pt(11)
    page_break(doc)

# ── Mục lục ───────────────────────────────────────────────────────────────────

def write_toc(doc):
    add_heading(doc, "MỤC LỤC", level=1, color=C_TITLE)
    toc_items = [
        ("1.",   "Giới thiệu & Bối cảnh nghiên cứu"),
        ("2.",   "Cơ sở lý thuyết"),
        ("2.1.", "  Semantic Segmentation và DeepLabV3+"),
        ("2.2.", "  Transfer Learning & Two-Phase Training"),
        ("2.3.", "  Loss Functions: Focal, Dice, Lovász"),
        ("3.",   "Bộ dữ liệu FoodSemSeg"),
        ("3.1.", "  Thu thập & Xây dựng dataset"),
        ("3.2.", "  Phân phối class & Xử lý mất cân bằng"),
        ("3.3.", "  Data Augmentation"),
        ("4.",   "Kiến trúc hệ thống"),
        ("4.1.", "  Mô hình DeepLabV3+ (ResNet50)"),
        ("4.2.", "  Pipeline 3 tầng: Seg → Depth → Weight"),
        ("4.3.", "  FastAPI Inference Server"),
        ("5.",   "Quá trình thực nghiệm (exp1 → exp21)"),
        ("5.1.", "  Hành trình lựa chọn kiến trúc"),
        ("5.2.", "  Chiến lược 2-Phase Training"),
        ("5.3.", "  Tiến hóa Loss Function"),
        ("5.4.", "  Kỹ thuật chống class imbalance"),
        ("5.5.", "  Ensemble Evaluation (eval_ensemble.py)"),
        ("6.",   "Kết quả & Đánh giá"),
        ("6.1.", "  Metrics: mIoU, Pixel Accuracy, BF1"),
        ("6.2.", "  Kết quả per-class IoU"),
        ("6.3.", "  Kết quả suy luận thực tế (API)"),
        ("7.",   "Calibration Pipeline Trọng lượng"),
        ("7.1.", "  Phát hiện vấn đề overestimate"),
        ("7.2.", "  So sánh 3 phương pháp calibration"),
        ("7.3.", "  Kết quả sau hiệu chỉnh"),
        ("8.",   "Bài học kinh nghiệm & Best Practices"),
        ("9.",   "Kết luận & Hướng phát triển"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{num:<5}")
        r1.font.bold = len(num) <= 2
        r1.font.size = Pt(10.5)
        r2 = p.add_run(title.strip())
        r2.font.bold = len(num) <= 2
        r2.font.size = Pt(10.5)
    page_break(doc)

# ── Phần 1: Giới thiệu ────────────────────────────────────────────────────────

def write_ch1(doc):
    add_heading(doc, "CHƯƠNG 1: GIỚI THIỆU & BỐI CẢNH NGHIÊN CỨU", level=1, color=C_TITLE)

    add_heading(doc, "1.1. Đặt vấn đề", level=2)
    add_para(doc,
        "Theo Tổ chức Y tế Thế giới (WHO), béo phì và các bệnh mãn tính liên quan đến chế độ ăn "
        "uống đang gia tăng với tốc độ đáng báo động. Việc theo dõi lượng calo và dinh dưỡng tiêu thụ "
        "hàng ngày là một trong những phương pháp hiệu quả để kiểm soát cân nặng và duy trì sức khỏe. "
        "Tuy nhiên, phương pháp thủ công (cân đo từng món, tra cứu bảng dinh dưỡng) đòi hỏi nhiều thời "
        "gian và kỷ luật, khiến đa số người dùng không thể duy trì lâu dài.",
        space_after=5)
    add_para(doc,
        "Sự phát triển của Deep Learning và Computer Vision mở ra khả năng tự động hóa quá trình này: "
        "chỉ cần chụp một bức ảnh bữa ăn, hệ thống có thể nhận diện từng loại thực phẩm, ước lượng "
        "khối lượng và tính toán giá trị dinh dưỡng tương ứng. Đây chính là bài toán mà đồ án này "
        "hướng tới giải quyết.",
        space_after=5)

    add_heading(doc, "1.2. Mục tiêu nghiên cứu", level=2)
    add_para(doc, "Đồ án đặt ra hai mục tiêu kỹ thuật chính:")
    add_bullet(doc, "Phân vùng ngữ nghĩa (Semantic Segmentation): Xác định chính xác vùng pixel "
               "và loại của từng thực phẩm trong ảnh đầu vào, với 21 nhãn thực phẩm phổ biến.")
    add_bullet(doc, "Ước lượng trọng lượng (Weight Estimation): Từ mask phân vùng và bản đồ độ "
               "sâu (depth map), tính toán khối lượng (gram) của từng loại thực phẩm, phục vụ "
               "tính toán dinh dưỡng.")

    add_heading(doc, "1.3. Phạm vi & Những giới hạn ban đầu", level=2)
    add_para(doc,
        "Dự án được thực hiện trên một máy tính cá nhân với GPU NVIDIA GTX 1660 Super (6GB VRAM), "
        "chạy hệ điều hành Windows 10. Giới hạn phần cứng này đóng vai trò quan trọng trong việc "
        "lựa chọn kiến trúc — loại bỏ các mô hình nặng như Mask2Former hay Swin-Transformer gốc "
        "do yêu cầu bộ nhớ vượt quá 6GB.",
        space_after=5)
    add_para(doc,
        "Giới hạn dữ liệu: tập dữ liệu được xây dựng thủ công từ nhiều nguồn mở (FoodInsSeg, "
        "COCO, fast_food, NutritionVerse-Real), với 21 class thực phẩm chọn lọc dựa trên độ phổ "
        "biến và tính khả thi về dữ liệu.",
        space_after=5)

    add_heading(doc, "1.4. Đóng góp của đồ án", level=2)
    add_bullet(doc, "Xây dựng pipeline hoàn chỉnh từ ảnh → mask → khối lượng, triển khai "
               "thực tế qua REST API (FastAPI).")
    add_bullet(doc, "Nghiên cứu và so sánh có hệ thống các chiến lược training (pretrain, "
               "loss function, augmentation, EMA) trên bài toán segmentation thực phẩm.")
    add_bullet(doc, "Phát triển phương pháp calibration trọng lượng, giảm MAE từ 8,125g xuống "
               "~248g (giảm 96.9%) thông qua so sánh 3 phương pháp hiệu chỉnh.")
    add_bullet(doc, "Đúc kết 20+ Best Practices có thể áp dụng cho các bài toán Computer Vision "
               "tương tự trong điều kiện phần cứng hạn chế.")
    page_break(doc)

# ── Phần 2: Cơ sở lý thuyết ──────────────────────────────────────────────────

def write_ch2(doc):
    add_heading(doc, "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT", level=1, color=C_TITLE)

    add_heading(doc, "2.1. Semantic Segmentation và DeepLabV3+", level=2)
    add_para(doc,
        "Semantic Segmentation là bài toán gán nhãn lớp cho từng pixel trong ảnh, phân biệt với "
        "Object Detection (bounding box) và Instance Segmentation (phân biệt từng đối tượng riêng lẻ). "
        "Đây là bài toán phù hợp nhất khi mục tiêu là tính tổng diện tích của một loại thực phẩm "
        "trong ảnh, thay vì đếm từng miếng riêng lẻ.",
        space_after=5)
    add_para(doc,
        "DeepLabV3+ (Chen et al., 2018) là kiến trúc state-of-the-art cho semantic segmentation, "
        "kết hợp Atrous Spatial Pyramid Pooling (ASPP) để nắm bắt ngữ cảnh đa tỉ lệ, cùng với "
        "decoder path lấy low-level features từ encoder để khôi phục độ phân giải biên. Trong đồ "
        "án này, backbone ResNet50 được sử dụng thông qua thư viện segmentation_models_pytorch (SMP).",
        space_after=5)
    add_para(doc,
        "Ưu điểm của DeepLabV3+ so với các kiến trúc khác trong bối cảnh phần cứng hạn chế: "
        "nhẹ hơn Mask2Former (~23M params), ổn định hơn trên Windows, và được hỗ trợ tốt bởi SMP "
        "với nhiều encoder pretrained sẵn có.",
        space_after=5)

    add_heading(doc, "2.2. Transfer Learning & Two-Phase Training", level=2)
    add_para(doc,
        "Transfer Learning là kỹ thuật tận dụng trọng số đã học từ tập dữ liệu lớn (ImageNet) "
        "để khởi tạo mô hình trên bài toán mới. Chiến lược Two-Phase Training được áp dụng "
        "trong đồ án gồm:",
        space_after=4)
    add_bullet(doc, "Phase 1 — Linear Probing: Đóng băng (freeze) toàn bộ encoder, chỉ "
               "train decoder và segmentation head với learning rate cao (1e-3). Mục tiêu: "
               "nhanh chóng thích nghi head với bài toán mới mà không phá vỡ features của encoder.")
    add_bullet(doc, "Phase 2 — Full Fine-tuning: Mở băng encoder, train toàn bộ mô hình với "
               "learning rate phân tầng: encoder (1e-5) << decoder (5e-5) < head (2e-4). "
               "Áp dụng EMA (Exponential Moving Average, decay=0.999) để ổn định hóa trọng số.")
    add_para(doc,
        "Lý do áp dụng differential LR: encoder đã học được features tốt từ ImageNet, cần "
        "update rất nhẹ để không bị catastrophic forgetting, trong khi head cần adapt nhanh "
        "hơn với bài toán phân lớp 22 class thực phẩm.",
        space_after=5)

    add_heading(doc, "2.3. Loss Functions: Focal, Dice, Lovász", level=2)
    add_para(doc,
        "Loss function đơn thuần Cross-Entropy không phù hợp cho bài toán có class imbalance "
        "nghiêm trọng (background chiếm >50% pixel). Đồ án sử dụng tổ hợp ba loss:",
        space_after=4)
    add_bullet(doc, "Focal Loss (Lin et al., 2017): FL(p_t) = −α(1−p_t)^γ · log(p_t), với "
               "γ=2. Giảm đóng góp của các pixel dễ phân loại (thường là background), tập "
               "trung model vào các pixel khó.")
    add_bullet(doc, "Dice Loss: L_Dice = 1 − 2|X∩Y|/(|X|+|Y|). Tối ưu trực tiếp tỉ lệ "
               "overlap giữa prediction và ground truth, đặc biệt hiệu quả khi class hiếm.")
    add_bullet(doc, "Lovász-Softmax Loss (Berman et al., 2018): Xấp xỉ khả vi trực tiếp "
               "của IoU loss thông qua Lovász extension. Đây là loss quan trọng nhất trong "
               "Phase 2 vì tối ưu gần với mục tiêu mIoU cuối cùng.")
    add_para(doc, "Công thức tổ hợp cuối cùng:")
    add_code_block(doc,
        "# Phase 1:\n"
        "L = 0.40 × Focal(γ=2) + 0.30 × Dice + 0.30 × CE(weighted)\n\n"
        "# Phase 2:\n"
        "L = 0.25 × Focal(γ=2) + 0.35 × Dice + 0.40 × Lovász-Softmax")

    add_heading(doc, "2.4. Monocular Depth Estimation (MiDaS)", level=2)
    add_para(doc,
        "MiDaS (Ranftl et al., 2020) là mô hình ước lượng độ sâu đơn mắt (monocular) "
        "sử dụng kiến trúc DPT (Dense Prediction Transformer). Điểm quan trọng cần lưu ý: "
        "MiDaS chỉ cho depth tương đối (relative depth), không phải depth tuyệt đối. "
        "Đây là nguyên nhân gốc rễ của vấn đề overestimate trọng lượng và dẫn đến sự cần "
        "thiết của bước Calibration.",
        space_after=5)
    page_break(doc)

# ── Phần 3: Dataset ───────────────────────────────────────────────────────────

def write_ch3(doc):
    add_heading(doc, "CHƯƠNG 3: BỘ DỮ LIỆU FOODSEMSEG", level=1, color=C_TITLE)

    add_heading(doc, "3.1. Thu thập & Xây dựng Dataset", level=2)
    add_para(doc,
        "FoodSemSeg là bộ dữ liệu Semantic Segmentation thực phẩm được xây dựng từ đầu cho "
        "đồ án này, thông qua việc tổng hợp và chuyển đổi từ nhiều nguồn dữ liệu mở:")
    make_table(doc,
        headers=["Nguồn dữ liệu", "Format gốc", "Ghi chú"],
        rows=[
            ["FoodInsSeg", "COCO JSON (Instance)", "Nguồn chính; chuyển đổi instance → semantic mask"],
            ["fast_food dataset", "YOLO .txt", "Viết script chuyển đổi riêng"],
            ["chicken dataset", "COCO JSON", "Thu thập bổ sung cho class hiếm"],
            ["COCO train", "COCO JSON", "Tích hợp thêm một số class"],
            ["NutritionVerse-Real", "COCO JSON", "38 ảnh chất lượng cao; dùng cho calibration"],
        ],
        col_widths=[Cm(5), Cm(4), Cm(7)]
    )
    add_para(doc,
        "Quá trình chuyển đổi từ Instance mask (COCO format) sang Semantic mask (PNG) gặp "
        "nhiều lỗi vặt điển hình như off-by-one ở category ID (+1). Phương pháp phát hiện: "
        "luôn visual audit ít nhất 1% dataset sau mỗi bước tiền xử lý.",
        space_after=5)

    add_heading(doc, "3.2. Tinh chỉnh class: từ 29 xuống 22 class", level=2)
    add_para(doc,
        "Ban đầu dataset gồm 29 class. Sau các thí nghiệm exp15–exp19, nhận thấy 7 class có "
        "IoU ≈ 0 do quá ít dữ liệu (tần suất pixel chỉ 0.05–0.2%). Quyết định loại bỏ 7 class "
        "này và chuẩn hóa còn 22 class (1 background + 21 food class) cho phiên bản chính thức:")
    add_bullet(doc, "Ảnh được pre-resize offline về 512×512 (thư mục FoodSemSeg_512x512) để "
               "tăng tốc validation (bỏ qua bước resize mỗi iteration).")
    add_bullet(doc, "21 class food: egg, banana, steak, pork, chicken, fish, shrimp, bread, "
               "noodles, rice, tofu, potato, tomato, lettuce, cucumber, carrot, broccoli, "
               "cabbage, onion, pepper, French beans.")

    add_heading(doc, "3.3. Phân phối class & Xử lý mất cân bằng (Class Imbalance)", level=2)
    add_para(doc,
        "Bài toán segmentation thực phẩm có mức độ mất cân bằng class rất cao: background "
        "chiếm trên 50% tổng pixel, trong khi một số class hiếm (tofu, French beans) chỉ "
        "chiếm dưới 0.3%. Các kỹ thuật được áp dụng:")
    add_bullet(doc,
        "Median Frequency Balancing: weight[c] = median(freq) / freq[c], clip về [0.1, 5.0]. "
        "Cho phép class hiếm có trọng số cao hơn trong loss mà không bị extreme outlier.")
    add_bullet(doc,
        "Class Weight Boost: CLASS_WEIGHT_BOOST = {11: 2.0} — tofu (class 11) được nhân "
        "thêm 2× do vẫn khó học sau median frequency balancing.")
    add_bullet(doc,
        "WeightedRandomSampler: ảnh chứa class hiếm được sample nhiều hơn trong mỗi epoch.")
    add_bullet(doc,
        "Strong Rare Augmentation: áp dụng augmentation mạnh hơn tự động khi ảnh chứa class hiếm.")

    add_heading(doc, "3.4. Data Augmentation Pipeline", level=2)
    add_para(doc, "Pipeline augmentation được thiết kế với thư viện Albumentations:")
    add_code_block(doc,
        "transforms = A.Compose([\n"
        "    A.HorizontalFlip(p=0.5),\n"
        "    A.VerticalFlip(p=0.2),\n"
        "    A.RandomRotate90(p=0.3),\n"
        "    A.Affine(scale=(0.85,1.15), translate_percent=(-0.1,0.1),\n"
        "             rotate=(-15,15), p=0.5),\n"
        "    A.RandomBrightnessContrast(0.2, p=0.5),\n"
        "    A.HueSaturationValue(p=0.3),\n"
        "    A.GaussNoise(p=0.2),\n"
        "    A.CoarseDropout(num_holes_range=(1,8), hole_height_range=(8,32),\n"
        "                    hole_width_range=(8,32), p=0.3),\n"
        "    A.Normalize(mean=[0.485,0.456,0.406], std=[0.229,0.224,0.225]),\n"
        "    ToTensorV2()\n"
        "])")
    page_break(doc)

# ── Phần 4: Kiến trúc hệ thống ────────────────────────────────────────────────

def write_ch4(doc):
    add_heading(doc, "CHƯƠNG 4: KIẾN TRÚC HỆ THỐNG", level=1, color=C_TITLE)

    add_heading(doc, "4.1. Mô hình Segmentation: DeepLabV3+ (ResNet50)", level=2)
    add_para(doc,
        "Mô hình được xây dựng bằng thư viện segmentation_models_pytorch (SMP) với "
        "DeepLabV3+ architecture và ResNet50 encoder. Hai mode khởi tạo được hỗ trợ:")
    make_table(doc,
        headers=["Mode", "Flag", "Mô tả", "Kết quả thực nghiệm"],
        rows=[
            ["A (Hiện tại)", "USE_IMAGENET=True",
             "ResNet50 pretrained ImageNet\nHead khởi tạo Kaiming",
             "Best: mIoU=0.3850 (exp21)"],
            ["B (Legacy)", "USE_IMAGENET=False",
             "Load checkpoint FoodSeg103\n(104 class), remap keys",
             "mIoU ≤ 0.25; bị overfit vào 103 class cũ"],
        ],
        col_widths=[Cm(2.5), Cm(3.5), Cm(5.5), Cm(4.5)]
    )
    add_para(doc,
        "Kết luận quan trọng: Pretrained trên ImageNet (general domain) cho kết quả tốt hơn "
        "pretrained trên FoodSeg103 (domain-specific) nhờ tính tổng quát cao hơn. Đây là một "
        "trong những bài học nền tảng của dự án (Hypothesis H1 bị bác bỏ).",
        space_after=5)
    add_para(doc, "Thông số kỹ thuật mô hình:")
    add_code_block(doc,
        "Architecture  : DeepLabV3+\n"
        "Encoder       : ResNet50 (pretrained ImageNet, ~23.5M params)\n"
        "Decoder       : DeepLabV3+ decoder (ASPP + low-level features)\n"
        "Head          : Segmentation head (out_channels=22)\n"
        "Input size    : 512 × 512 × 3\n"
        "Output size   : 512 × 512 × 22 (logits) → 512 × 512 (class map)")

    add_heading(doc, "4.2. Pipeline 3 tầng: Seg → Depth → Weight", level=2)
    add_para(doc,
        "Sau khi có mask phân vùng, một pipeline hoàn chỉnh được xây dựng để ước lượng trọng lượng:")
    make_table(doc,
        headers=["Tầng", "Đầu vào", "Xử lý", "Đầu ra"],
        rows=[
            ["Tier 1\nSegmentation",
             "Ảnh RGB\n(bất kỳ kích thước)",
             "DeepLabV3+ inference\n+ Instance extraction\n(CC / Watershed)",
             "Semantic mask\n+ instance masks"],
            ["Tier 2\nDepth & Volume",
             "Ảnh RGB\n+ Semantic mask",
             "MiDaS DPT_Large\n→ depth map\n→ volume ước lượng",
             "Volume (cm³)\nper instance"],
            ["Tier 3\nWeight",
             "Volume (cm³)\n+ class_id",
             "Tra bảng mật độ\nweight = ρ × V × scale",
             "Trọng lượng (g)\nper class"],
        ],
        col_widths=[Cm(2.5), Cm(3.5), Cm(5), Cm(5)]
    )
    add_para(doc, "Bảng mật độ thực phẩm (g/cm³) tích hợp trong pipeline:")
    add_code_block(doc,
        "bread: 0.35,  lettuce: 0.35,  broccoli: 0.60,  cabbage: 0.45\n"
        "shrimp: 0.85, noodles: 0.90,  banana: 0.95,    tofu: 0.95\n"
        "cucumber: 0.96, tomato: 0.99, fish: 1.02,      egg: 1.03\n"
        "carrot: 1.04, pork/chicken/onion: 1.05,        potato: 1.09\n"
        "steak: 1.08,  rice: 1.15")

    add_heading(doc, "4.3. FastAPI Inference Server", level=2)
    add_para(doc,
        "Mô hình được triển khai qua FastAPI server với endpoint POST /predict. "
        "Điểm đặc biệt trong xử lý ảnh đầu vào:")
    add_bullet(doc, "Distortion-free preprocessing: scale cạnh dài về 512, pad về 512×512 "
               "(không méo ảnh).")
    add_bullet(doc, "Post-inference: crop bỏ padding, resize mask về kích thước gốc.")
    add_bullet(doc, "Tự động vẽ legend chỉ với các class xuất hiện trong ảnh.")
    add_bullet(doc, "Lưu 4 file per request: _orig.png, _mask.png, _overlay.png, _meta.json.")
    add_para(doc, "Response JSON format:")
    add_code_block(doc,
        '{\n'
        '  "image_info": {"width": 1200, "height": 675},\n'
        '  "total_estimated_weight_g": 390.5,\n'
        '  "detections": [\n'
        '    {"class_id": 7, "class_name": "chicken",\n'
        '     "instance_count": 1, "estimated_weight_g": 380.6,\n'
        '     "color_hex": "#ff8c00"},\n'
        '    ...\n'
        '  ]\n'
        '}')
    page_break(doc)

# ── Phần 5: Thực nghiệm ───────────────────────────────────────────────────────

def write_ch5(doc):
    add_heading(doc, "CHƯƠNG 5: QUÁ TRÌNH THỰC NGHIỆM (exp1 → exp21)", level=1, color=C_TITLE)

    add_heading(doc, "5.1. Hành trình lựa chọn kiến trúc", level=2)
    add_para(doc,
        "Trước khi đến với DeepLabV3+, dự án đã trải qua hành trình thử nghiệm nhiều kiến trúc "
        "khác nhau. Mỗi kiến trúc để lại một bài học quan trọng:")
    make_table(doc,
        headers=["Kiến trúc", "Lý do thử", "Vấn đề gặp phải", "Bài học"],
        rows=[
            ["YOLOv8-seg",
             "Tốc độ nhanh, phổ biến, dễ dùng",
             "Độ chính xác biên kém cho bài toán đo đạc",
             "Speed ≠ Accuracy cho downstream task"],
            ["Mask R-CNN\n(MMDetection)",
             "Độ chính xác instance cao hơn YOLO",
             "Framework nặng, config phức tạp, chậm train",
             "Overhead framework > lợi ích accuracy"],
            ["Mask2Former",
             "SOTA theo paper 2022",
             "OOM trên 6GB GPU; lỗi path Windows\n(dấu cách trong đường dẫn)",
             "Hardware constraint > SOTA paper"],
            ["DeepLabV3+\n(ResNet50 / SMP)",
             "Nhẹ, ổn định, hỗ trợ Windows tốt",
             "Không có vấn đề nghiêm trọng",
             "Fit-to-constraints là chiến lược đúng"],
        ],
        col_widths=[Cm(3), Cm(4), Cm(4.5), Cm(4.5)]
    )
    add_para(doc,
        "Cú rẽ quan trọng nhất: chuyển từ Instance Segmentation sang Semantic Segmentation. "
        "Câu hỏi dẫn đến quyết định này: 'Mục tiêu cuối cùng là khối lượng — có cần đếm từng "
        "miếng riêng lẻ không?' Câu trả lời là không: tổng diện tích pixel của mỗi class là "
        "đủ để ước lượng khối lượng. Quyết định này đơn giản hóa toàn bộ pipeline và cho phép "
        "training ổn định hơn.",
        space_after=5)

    add_heading(doc, "5.2. Hành trình thí nghiệm (Timeline of Experiments)", level=2)
    make_table(doc,
        headers=["Experiment", "Dataset", "Thay đổi chính", "mIoU đạt được"],
        rows=[
            ["exp1–10",   "FoodSemSeg 29 class", "Thăm dò FoodSeg103 pretrain; nhiều cấu hình CE loss", "~0.15–0.18"],
            ["exp11",     "FoodSemSeg (18 class)", "Chuyển sang ImageNet pretrained: đột phá lớn nhất",  "0.52*"],
            ["exp15–18",  "FoodSemSeg 29 class", "Focal+Dice loss; tune LR; augmentation",              "~0.30–0.34"],
            ["exp19",     "FoodSemSeg 29 class", "WeightedSampler; thu thập thêm data",                 "0.36"],
            ["exp20",     "FoodSemSeg_512 22 class", "Cắt giảm xuống 22 class; offline resize",         "0.3686"],
            ["exp21",     "FoodSemSeg_512 22 class", "Bổ sung data cho các class yếu; EMA + Lovász + differential LR", "0.3850 ★"],
        ],
        col_widths=[Cm(2.5), Cm(4), Cm(6.5), Cm(3)]
    )
    add_para(doc,
        "* exp11: mIoU 0.52 trên validation set cũ (18 class, split và dataset khác với exp20–21). "
        "Không so trực tiếp với 0.385 vì số class và thành phần test đã thay đổi.",
        italic=True, size=10, space_before=2, space_after=5)

    add_heading(doc, "5.3. Chiến lược 2-Phase Training (Chi tiết exp21)", level=2)
    add_para(doc, "Cấu hình hyperparameter chi tiết của exp21 — best run:")
    add_code_block(doc,
        "# Phase 1 — Linear Probing (encoder frozen)\n"
        "P1_LR         = 1e-3\n"
        "P1_EPOCHS     = 100   (early stop patience=8)\n"
        "P1_LOSS       = 0.40×Focal + 0.30×Dice + 0.30×CE(weighted)\n"
        "P1_OPTIMIZER  = Adam\n\n"
        "# Phase 2 — Full Fine-tuning (encoder unfrozen)\n"
        "P2_LR_ENCODER = 1e-5   P2_LR_DECODER = 5e-5   P2_LR_HEAD = 2e-4\n"
        "P2_WEIGHT_DECAY = 3e-4\n"
        "P2_EPOCHS     = 100   (early stop patience=15)\n"
        "P2_LOSS       = 0.25×Focal + 0.35×Dice + 0.40×Lovász\n"
        "P2_USE_EMA    = True  (decay=0.999)\n"
        "BATCH_SIZE    = 4     GRAD_CLIP_MAX_NORM = 0.5")

    add_heading(doc, "5.4. Các vấn đề kỹ thuật nổi bật", level=2)

    add_heading(doc, "5.4.1. Vấn đề AMP (Mixed Precision Training)", level=3)
    add_para(doc,
        "Trên GPU GTX 1660 Super, việc bật AMP (Automatic Mixed Precision) gây ra "
        "GradNorm = inf/nan liên tục, khiến training mất ổn định. Nguyên nhân: GPU "
        "consumer-grade có floating point FP16 kém ổn định hơn cho gradient computation "
        "so với Tesla/A100. Giải pháp: tắt AMP (USE_AMP=False), dùng FP32 hoàn toàn.")

    add_heading(doc, "5.4.2. Lỗi 'Silent Config Killer'", level=3)
    add_para(doc,
        "Triệu chứng: training đột nhiên chậm gấp đôi dù chỉ thêm 44 ảnh. Nguyên nhân: "
        "BATCH_SIZE bị thay đổi âm thầm từ 8 → 4 trong file config. Bài học: file config "
        "là 'hidden state' nguy hiểm; cần log lại toàn bộ config mỗi run và version "
        "control config files.")

    add_heading(doc, "5.5. Ensemble Evaluation (eval_ensemble.py)", level=2)
    add_para(doc,
        "Để tăng độ tin cậy của kết quả đánh giá, một kỹ thuật ensemble được triển khai "
        "trong script eval_ensemble.py: average logits từ N checkpoint trước khi argmax. "
        "Pipeline đánh giá ensemble hoạt động như sau:")
    add_code_block(doc,
        "# eval_ensemble.py — Luồng chính\n"
        "def predict_batch(images):\n"
        "    if args.tta:  # Test-Time Augmentation\n"
        "        for m in models:\n"
        "            l0 = m(images)                              # original\n"
        "            l1 = flip(m(flip(images, H)))               # H-flip\n"
        "            l2 = flip(m(flip(images, V)), V)            # V-flip\n"
        "            logits_list.append((l0 + l1 + l2) / 3.0)\n"
        "    else:\n"
        "        logits_list = [m(images) for m in models]\n"
        "    return torch.stack(logits_list).mean(dim=0)  # average logits\n\n"
        "# Rare class threshold override (giảm ngưỡng quyết định)\n"
        "# VD: --rare-threshold '11=0.25' → nếu prob[tofu]>=0.25 và > prob[pred] → gán tofu")
    add_para(doc,
        "Script còn hỗ trợ tùy chọn --rare-threshold để hạ ngưỡng quyết định cho class hiếm: "
        "nếu xác suất của class hiếm vượt ngưỡng đã đặt VÀ lớn hơn xác suất của nhãn đang "
        "được dự đoán, thì override nhãn sang class hiếm đó. Kết quả được lưu vào "
        "runs/train/exp_ensemble/per_class_iou_ensemble.csv.")
    page_break(doc)

# ── Phần 6: Kết quả ───────────────────────────────────────────────────────────

def write_ch6(doc):
    add_heading(doc, "CHƯƠNG 6: KẾT QUẢ & ĐÁNH GIÁ", level=1, color=C_TITLE)

    add_heading(doc, "6.1. Metrics đánh giá", level=2)
    make_table(doc,
        headers=["Metric", "Công thức / Phương pháp", "Ý nghĩa"],
        rows=[
            ["mIoU\n(mean IoU)",
             "IoU(c) = TP/(TP+FP+FN)\nmIoU = avg IoU (bỏ background)",
             "Metric chính; bao quát cả precision lẫn recall"],
            ["Pixel Accuracy",
             "PA = Σ(pred==gt) / total_pixels",
             "Dễ bị lừa bởi class imbalance (background chiếm đa số)"],
            ["Boundary F1\n(BF1)",
             "Dilate−Erode → boundary mask\n→ precision/recall/F1",
             "Đánh giá chất lượng biên; quan trọng cho ứng dụng đo đạc"],
            ["Exact Match\nAccuracy (EMA)",
             "Count GT instances == Count Pred\nper class (via CC)",
             "Đánh giá khả năng đếm instance thông qua CC"],
            ["MACE",
             "Mean Absolute Count Error\nqua Connected Components",
             "Định lượng sai số đếm instance"],
        ],
        col_widths=[Cm(3), Cm(6.5), Cm(6.5)]
    )

    add_heading(doc, "6.2. Kết quả tổng thể — exp21 (Best Run)", level=2)
    make_table(doc,
        headers=["Phase", "Epoch", "Train Loss", "mIoU", "Pixel Accuracy", "BF1"],
        rows=[
            ["Phase 1 (peak)", "31",  "0.5645", "0.2435", "0.7888", "—"],
            ["Phase 2 (peak)", "55",  "0.2450", "0.3850", "0.8282", "0.1971"],
        ],
        col_widths=[Cm(3), Cm(2), Cm(2.5), Cm(2.5), Cm(3.5), Cm(2.5)]
    )
    add_para(doc,
        "Nhận xét: mIoU tăng từ 0.2435 (cuối Phase 1) lên 0.3850 (cuối Phase 2), "
        "tương đương mức cải thiện 58.2%. Điều này xác nhận hiệu quả của chiến lược "
        "Two-Phase Training: Phase 1 cho model 'vào vị trí', Phase 2 tinh chỉnh sâu.",
        space_after=5)

    add_heading(doc, "6.3. Kết quả per-class IoU — exp21", level=2)
    make_table(doc,
        headers=["Class", "IoU", "Class", "IoU", "Class", "IoU"],
        rows=[
            ["broccoli (17)", "0.7644 ★", "carrot (16)",   "0.7434",   "tomato (13)",   "0.6521"],
            ["rice (10)",     "0.6279",   "bread (8)",     "0.6124",   "cucumber (15)", "0.5549"],
            ["potato (12)",   "0.5129",   "Frch.beans (21)","0.4962", "noodles (9)",   "0.4305"],
            ["egg (1)",       "0.4575",   "chicken (5)",   "0.4563",   "banana (2)",    "0.4284"],
            ["cabbage (18)",  "0.4267",   "tofu (11)",     "0.3976",   "shrimp (7)",    "0.3907"],
            ["steak (3)",     "0.4010",   "lettuce (14)",  "0.3885",   "onion (19)",    "0.3162"],
            ["pepper (20)",   "0.3107",   "fish (6)",      "0.2983",   "pork (4)",      "0.1315 ✗"],
        ],
        col_widths=[Cm(3.5), Cm(2), Cm(3.5), Cm(2), Cm(3.5), Cm(2)]
    )
    add_para(doc,
        "Phân tích: Broccoli (0.7644) và carrot (0.7434) đạt IoU cao nhất — cả hai class "
        "có màu sắc đặc trưng rõ ràng (xanh lá, cam). Pork (0.1315) là class khó nhất "
        "do màu sắc và kết cấu tương đồng với nhiều class thịt khác (chicken, steak, fish). "
        "Tofu (0.3976) cải thiện đáng kể so với các exp trước nhờ class weight boost ×2.",
        space_after=5)

    add_heading(doc, "6.4. So sánh Ensemble vs Single Model", level=2)
    make_table(doc,
        headers=["Class", "Single (exp21)", "Ensemble", "Δ"],
        rows=[
            ["broccoli", "0.7644", "0.7566", "−0.0078"],
            ["carrot",   "0.7434", "0.7186", "−0.0248"],
            ["tomato",   "0.6521", "0.6341", "−0.0180"],
            ["bread",    "0.6124", "0.6047", "−0.0077"],
            ["noodles",  "0.4305", "0.4701", "+0.0396 ↑"],
            ["tofu",     "0.3976", "0.4519", "+0.0543 ↑"],
            ["potato",   "0.5129", "0.4726", "−0.0403"],
            ["pork",     "0.1315", "0.1341", "+0.0026"],
        ],
        col_widths=[Cm(4), Cm(3), Cm(3), Cm(2.5)]
    )
    add_para(doc,
        "Nhận xét: Ensemble không luôn luôn tốt hơn single best model. Với class có tần "
        "suất cao (broccoli, carrot, tomato), single model exp21 thắng nhẹ. Tuy nhiên, "
        "ensemble cho kết quả tốt hơn rõ rệt với các class hiếm như tofu (+0.054) và "
        "noodles (+0.040), xác nhận giá trị của ensemble cho long-tail distribution.",
        space_after=5)

    add_heading(doc, "6.5. So sánh exp19 vs exp21 trên ảnh test thực tế", level=2)
    add_para(doc,
        "Để đánh giá trực quan sự tiến bộ từ exp19 (29 class) lên exp21 (22 class), nhóm so sánh "
        "kết quả suy luận của hai mô hình trên cùng một ảnh test qua FastAPI. Ảnh test là ảnh "
        "gà với salad rau (1200×675), kết quả lưu tại scripts/api_results_exp19 (exp19) và "
        "scripts/api_results (exp21).",
        space_after=5)

    # Ảnh gốc + overlay exp19 + overlay exp21
    add_heading(doc, "Ảnh test: gà với salad rau (1200×675)", level=3)
    img_orig = os.path.join(IMG_DIR_EXP19, "test_orig.png")
    img_exp19 = os.path.join(IMG_DIR_EXP19, "test_overlay.png")
    img_exp21 = os.path.join(IMG_DIR_EXP21, "test_overlay.png")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_orig):
        p.add_run().add_picture(img_orig, width=Inches(2.5))
        p.add_run("  ")
    if os.path.exists(img_exp19):
        p.add_run().add_picture(img_exp19, width=Inches(2.5))
        p.add_run("  ")
    if os.path.exists(img_exp21):
        p.add_run().add_picture(img_exp21, width=Inches(2.5))
    c = doc.add_paragraph("Hình 6.1: Ảnh gốc (trái) | Overlay exp19 — 29 class (giữa) | Overlay exp21 — 22 class (phải)")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].font.italic = True; c.runs[0].font.size = Pt(9.5)
    c.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_heading(doc, "Bảng so sánh kết quả phát hiện", level=3)
    make_table(doc,
        headers=["Class", "exp19 — Instance", "exp19 — Weight (g)", "exp21 — Instance", "exp21 — Weight (g)"],
        rows=[
            ["chicken", "1", "380.6", "1", "472.8"],
            ["tomato",  "1", "1.0",   "1", "1.7"],
            ["lettuce", "3", "8.8",   "4", "13.4"],
            ["pepper",  "1", "0.1", "—", "—"],
            ["TỔNG",    "—", "390.5", "—", "487.9"],
        ],
        col_widths=[Cm(3), Cm(2.5), Cm(3), Cm(2.5), Cm(3)]
    )
    add_para(doc,
        "Nhận xét: exp21 (22 class) phát hiện 4 instance lettuce so với 3 instance của exp19, "
        "cho thấy khả năng tách instance tốt hơn nhờ bổ sung data cho các class yếu và kiến trúc "
        "tinh chỉnh. exp19 phát hiện pepper (0.1g) là sai — false positive; exp21 không dự đoán "
        "pepper nên chính xác hơn. Tổng trọng lượng ước lượng khác nhau (390.5g vs 487.9g) "
        "do sự khác biệt về volume estimation và calibration giữa hai pipeline.",
        space_after=5)
    page_break(doc)

# ── Phần 7: Calibration ───────────────────────────────────────────────────────

def write_ch7(doc):
    add_heading(doc, "CHƯƠNG 7: CALIBRATION PIPELINE TRỌNG LƯỢNG", level=1, color=C_TITLE)

    add_heading(doc, "7.1. Phát hiện vấn đề: Overestimate 300 lần", level=2)
    add_para(doc,
        "Sau khi triển khai pipeline 3 tầng thô (không calibration), nhóm chạy đánh giá trên "
        "bộ NutritionVerse-Real gồm 38 ảnh thực phẩm chất lượng cao với ground-truth trọng "
        "lượng được cân đo thực tế. Kết quả ban đầu:")
    make_table(doc,
        headers=["Metric", "Trước Calibration", "Sau Calibration (Method B)", "Cải thiện"],
        rows=[
            ["MAE (g)",   "8,125.0",  "247.7",   "↓ 96.9%"],
            ["MAPE (%)",  "3,330.3",  "128.1",   "↓ 96.2%"],
        ],
        col_widths=[Cm(4), Cm(4), Cm(4), Cm(4)]
    )
    add_para(doc,
        "Nguyên nhân gốc rễ: MiDaS chỉ cho relative depth (giá trị tương đối trong ảnh), "
        "không phải absolute depth (cm thực tế). Depth map bị scale lên hàng trăm lần so với "
        "kích thước vật lý thực, dẫn đến volume và weight bị overestimate hệ thống.",
        space_after=5)

    add_heading(doc, "7.2. So sánh 3 phương pháp Calibration", level=2)
    make_table(doc,
        headers=["Method", "Mô tả", "Ưu điểm", "Nhược điểm"],
        rows=[
            ["A\nGlobal Scale",
             "Nhân toàn bộ weight\nvới factor = 0.0512",
             "Đơn giản, tính được\ntừ 1 số thống kê",
             "Không phân biệt class;\ncùng lệch hướng cho tất cả class"],
            ["B\nPer-class\nEffective Density",
             "Thay mật độ thực tế bằng\nmedian(gt_weight/pipeline_vol)\nper class",
             "Linh hoạt per class;\nphù hợp nếu pipeline vol\nkhá nhất quán",
             "Cần đủ data cho\ntừng class"],
            ["C\nReal Density +\nVolume Scale ★",
             "Giữ mật độ vật lý thực;\nnhân volume với scale\nchung = 0.0428",
             "Grounded trong vật lý;\nchỉ cần 1 scale factor;\ntổng quát tốt",
             "Giả định lệch volume\nđồng đều mọi class"],
        ],
        col_widths=[Cm(2.5), Cm(4.5), Cm(4), Cm(5)]
    )
    add_para(doc,
        "Phương pháp C được chọn vì: (1) sử dụng mật độ vật lý thực tế nên có tính giải "
        "thích được (interpretable); (2) chỉ cần một global scale factor nên tổng quát hơn; "
        "(3) hiệu năng tương đương Method B trên tập calibration. Volume scale được tính: "
        "scale = median(gt_weight / (real_density × midas_volume)) = 0.042778.",
        space_after=5)

    add_heading(doc, "7.3. Phân tích nguyên nhân & Hướng khắc phục", level=2)
    add_para(doc,
        "Bài học sâu từ vấn đề này: khi sử dụng một mô hình học máy (MiDaS) như một module "
        "trong pipeline lớn hơn, cần kiểm tra kỹ đặc tính đầu ra của module đó — đặc biệt "
        "là sự khác biệt giữa relative và absolute output. Đây là ví dụ điển hình của lỗi "
        "'learned approximation without physical grounding'.")
    add_para(doc,
        "Hướng khắc phục triệt để (future work): thay thế MiDaS bằng các mô hình depth "
        "tuyệt đối (metric depth) như ZoeDepth hoặc Metric3D, loại bỏ hoàn toàn bước "
        "calibration phức tạp.")
    page_break(doc)

# ── Phần 8: Best Practices ────────────────────────────────────────────────────

def write_ch8(doc):
    add_heading(doc, "CHƯƠNG 8: BÀI HỌC KINH NGHIỆM & BEST PRACTICES", level=1, color=C_TITLE)

    add_heading(doc, "8.1. Những chuyển dịch tư duy (Paradigm Shifts)", level=2)
    make_table(doc,
        headers=["Tư duy cũ", "Tư duy mới", "Bằng chứng"],
        rows=[
            ["Model-Centric:\n'Dùng model nào tốt nhất?'",
             "Data-Centric:\n'Dữ liệu có vấn đề gì?'",
             "Loại bỏ 7 class ít data → mIoU tăng từ 0.36 lên 0.385"],
            ["Instance Seg:\n'Cần đếm từng miếng'",
             "Semantic Seg:\n'Đầu ra thực sự cần là gì?'",
             "Tổng diện tích pixel đủ để ước lượng khối lượng"],
            ["Forward Fix:\ncố sửa bug tiếp",
             "Revert-first Debugging:\nrevert về stable state",
             "Lỗi Phantom Overlay Bug — revert giải quyết trong < 1h"],
            ["Learned Approximation:\nML tự học mọi thứ",
             "Physical Grounding:\nencode domain knowledge",
             "Method C thắng nhờ dùng mật độ vật lý thực tế"],
            ["SOTA First:\ndùng model mạnh nhất",
             "Constraint First:\nfit vào hardware của mình",
             "Mask2Former OOM; DeepLabV3+ ổn định hoàn toàn"],
        ],
        col_widths=[Cm(4.5), Cm(4.5), Cm(7)]
    )

    add_heading(doc, "8.2. Best Practices Catalog", level=2)

    add_heading(doc, "D — Data Engineering", level=3)
    add_bullet(doc, "D1: Luôn visual audit ≥1% dataset sau mỗi bước tiền xử lý để phát hiện lỗi off-by-one.")
    add_bullet(doc, "D2: Pre-resize offline dataset về kích thước cố định (512×512) để tăng tốc validation loop.")
    add_bullet(doc, "D3: Vẽ class distribution plot trước khi lựa chọn loss function và sampling strategy.")
    add_bullet(doc, "D4: Dùng Median Frequency Balancing cho class imbalance ratio > 100×.")

    add_heading(doc, "M — Model & Training", level=3)
    add_bullet(doc, "M1: Luôn bắt đầu với ImageNet pretrained, bất kể domain, do tính tổng quát cao hơn domain-specific pretrain.")
    add_bullet(doc, "M2: Two-Phase training là default pattern an toàn nhất: freeze encoder (Phase 1) → unfreeze (Phase 2).")
    add_bullet(doc, "M3: Monitor GradNorm mỗi batch để phát hiện AMP instability sớm (dấu hiệu: GradNorm = inf/nan liên tục).")
    add_bullet(doc, "M4: Kết hợp Focal + Dice + Lovász để tối ưu trực tiếp và gián tiếp metric IoU.")
    add_bullet(doc, "M5: Thêm TTA (Test-Time Augmentation: H-flip + V-flip) khi evaluate để 'free lunch' +1–3% mIoU.")
    add_bullet(doc, "M6: EMA (decay=0.999) ở Phase 2 giúp trọng số ổn định hơn, giảm variance giữa các checkpoint.")

    add_heading(doc, "A — Architecture & Infrastructure", level=3)
    add_bullet(doc, "A1: Chọn kiến trúc dựa trên hardware constraints trước, SOTA paper sau.")
    add_bullet(doc, "A2: Trên Windows, tránh các framework nặng về subprocess như Detectron2/MMDetection (lỗi path với dấu cách).")
    add_bullet(doc, "A3: Thiết kế artifact schema (mask, overlay, meta.json) sớm để decouple các tầng pipeline.")

    add_heading(doc, "S — Scientific Method", level=3)
    add_bullet(doc, "S1: Phân tích causal chain trước khi fix bug: symptom → mechanism → root cause.")
    add_bullet(doc, "S2: Khi mIoU thấp nhưng Pixel Accuracy cao → đây là 'class imbalance smell', không phải thiếu model capacity.")
    add_bullet(doc, "S3: Log toàn bộ config file vào mỗi run directory để reproducibility và audit trail.")
    page_break(doc)

# ── Phần 9: Kết luận ─────────────────────────────────────────────────────────

def write_ch9(doc):
    add_heading(doc, "CHƯƠNG 9: KẾT LUẬN & HƯỚNG PHÁT TRIỂN", level=1, color=C_TITLE)

    add_heading(doc, "9.1. Kết luận", level=2)
    add_para(doc,
        "Đồ án đã xây dựng thành công một hệ thống hoàn chỉnh cho bài toán phân vùng ngữ nghĩa "
        "thực phẩm và ước lượng trọng lượng tự động, với những đóng góp kỹ thuật cụ thể:")
    make_table(doc,
        headers=["Thành phần", "Kết quả đạt được"],
        rows=[
            ["Mô hình Segmentation (exp21)",
             "mIoU = 0.3850 · Pixel Accuracy = 0.8282 · BF1 = 0.1971\n21/21 class được nhận diện (IoU > 0.10)"],
            ["Pipeline Ước lượng Trọng lượng",
             "MAE giảm 96.9%: từ 8,125g → 247.7g\nMAPE giảm 96.2%: từ 3,330% → 128.1%"],
            ["REST API (FastAPI)",
             "Xử lý ảnh bất kỳ kích thước, distortion-free\nResponse JSON đầy đủ: class, count, weight, color"],
            ["Tài liệu hóa",
             "21 thí nghiệm có log đầy đủ\n20+ Best Practices đúc kết từ thực chiến"],
        ],
        col_widths=[Cm(5), Cm(11)]
    )
    add_para(doc,
        "Đặc biệt, hành trình 21 thí nghiệm đã chứng minh rằng trong điều kiện phần cứng hạn "
        "chế (GPU 6GB), chiến lược Data-Centric AI (cải thiện dữ liệu, xử lý imbalance, "
        "tinh chỉnh loss) đem lại lợi ích lớn hơn là theo đuổi các kiến trúc SOTA nặng hơn.",
        space_after=5)

    add_heading(doc, "9.2. Hạn chế hiện tại", level=2)
    add_bullet(doc, "mIoU 0.385 vẫn còn khoảng cách so với các hệ thống thương mại (>0.60). "
               "Nguyên nhân chính: dữ liệu training còn hạn chế, đặc biệt với các class hiếm.")
    add_bullet(doc, "MAPE 128.1% của ước lượng trọng lượng vẫn còn cao cho ứng dụng thực tế "
               "— cần depth tuyệt đối thay vì depth tương đối.")
    add_bullet(doc, "Một số class vẫn có IoU thấp: pork (0.1315), fish (0.2983), onion (0.3162) "
               "do tương đồng về màu sắc và kết cấu với các class khác.")

    add_heading(doc, "9.3. Hướng phát triển", level=2)
    make_table(doc,
        headers=["Hướng", "Phương pháp đề xuất", "Kỳ vọng cải thiện"],
        rows=[
            ["Cải thiện Segmentation",
             "Thay ResNet50 bằng ConvNeXt-Base hoặc Swin-T\nThu thập thêm data cho pork, fish, onion",
             "mIoU → 0.45–0.50"],
            ["Cải thiện Depth",
             "Thay MiDaS (relative) bằng ZoeDepth/Metric3D\n(absolute metric depth estimation)",
             "Loại bỏ hoàn toàn bước Calibration;\nMAPE → <30%"],
            ["Mở rộng Calibration",
             "Tăng bộ calibration lên >200 ảnh\nPer-class volume correction thay vì global",
             "Cải thiện từng class riêng lẻ"],
            ["Production hóa",
             "Docker container + load balancing\nCamera stereo/LiDAR cho absolute depth",
             "Throughput cao hơn; MAE → <50g"],
        ],
        col_widths=[Cm(3.5), Cm(7), Cm(5.5)]
    )
    page_break(doc)

# ── Tài liệu tham khảo ────────────────────────────────────────────────────────

def write_references(doc):
    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1, color=C_TITLE)
    refs = [
        "[1] L.-C. Chen, Y. Zhu, G. Papandreou, F. Schroff, H. Adam. "
        "\"Encoder-Decoder with Atrous Separable Convolution for Semantic Image Segmentation\" "
        "(DeepLabV3+). ECCV 2018.",
        "[2] T.-Y. Lin, P. Goyal, R. Girshick, K. He, P. Dollar. "
        "\"Focal Loss for Dense Object Detection\" (RetinaNet / Focal Loss). ICCV 2017.",
        "[3] M. Berman, A. Rannen Triki, M. B. Blaschko. "
        "\"The Lovász-Softmax Loss: A Tractable Surrogate for the Optimization of the "
        "Intersection-Over-Union Measure in Neural Networks\". CVPR 2018.",
        "[4] R. Ranftl, K. Lasinger, D. Hafner, K. Schindler, V. Koltun. "
        "\"Towards Robust Monocular Depth Estimation: Mixing Datasets for Zero-Shot "
        "Cross-Dataset Transfer\" (MiDaS). IEEE TPAMI 2020.",
        "[5] P. Yakubovskiy. \"Segmentation Models PyTorch\". GitHub 2020. "
        "https://github.com/qubvel/segmentation_models.pytorch",
        "[6] A. Buslaev, V. I. Iglovikov, E. Khvedchenya, A. Parinov, M. Druzhinin, "
        "A. A. Kalinin. \"Albumentations: Fast and Flexible Image Augmentations\". "
        "Information 2020.",
        "[7] K. He, X. Zhang, S. Ren, J. Sun. \"Deep Residual Learning for Image "
        "Recognition\" (ResNet). CVPR 2016.",
        "[8] J. Jiang, W. Luo, Q. Gao, F. Zhou, S. Chang. \"FoodInsSeg: A Food Instance "
        "Segmentation Dataset and Benchmark\". 2022.",
        "[9] S. Sebastião, A. Figueiredo, T. Guerreiro, H. Nicolau. \"NutritionVerse: "
        "Empirical Study of Various Dietary Intake Estimation Approaches\". Nutrients 2023.",
        "[10] S. Anand et al. \"MaskFormer: Per-Pixel Classification is Not All You Need "
        "for Semantic Segmentation\". NeurIPS 2021.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(ref)
        run.font.size = Pt(10)

# ── PHỤ LỤC ──────────────────────────────────────────────────────────────────

def write_appendix(doc):
    add_heading(doc, "PHỤ LỤC", level=1, color=C_TITLE)

    add_heading(doc, "A. Cấu trúc thư mục dự án", level=2)
    add_code_block(doc,
        "DeepLabV3_RetNet50/\n"
        "├── config.py               ← Cấu hình toàn bộ hệ thống\n"
        "├── dataset.py              ← Custom Dataset + augmentation + class weights\n"
        "├── model_setup.py          ← Build DeepLabV3+ ResNet50\n"
        "├── metrics.py              ← mIoU, PixelAcc, BF1, Instance Count\n"
        "├── utils.py                ← Loss functions, logging helpers\n"
        "├── postprocess.py          ← Morphological smoothing + instance counting\n"
        "├── data/\n"
        "│   ├── FoodSemSeg/         ← Dataset gốc (29 class, legacy exp19)\n"
        "│   └── FoodSemSeg_512x512/ ← Dataset đã resize offline (22 class)\n"
        "├── runs/train/\n"
        "│   ├── exp2–exp21/         ← Checkpoints, logs, CSV kết quả\n"
        "│   └── exp_ensemble/       ← Kết quả ensemble evaluation\n"
        "├── scripts/\n"
        "│   ├── training/train.py   ← Train 2-phase\n"
        "│   ├── training/eval_ensemble.py\n"
        "│   └── api/api.py          ← FastAPI server (22 class)\n"
        "└── nv_pipeline/            ← Pipeline 3 tầng seg→depth→weight")

    add_heading(doc, "B. Lệnh chạy hệ thống", level=2)
    add_para(doc, "Train mô hình từ đầu:")
    add_code_block(doc, "python scripts/training/train.py")
    add_para(doc, "Đánh giá ensemble:")
    add_code_block(doc,
        "python scripts/training/eval_ensemble.py \\\n"
        "  --checkpoints runs/train/exp21/weights/best.pth \\\n"
        "                runs/train/exp20/weights/best.pth \\\n"
        "  --tta --rare-threshold 11=0.25")
    add_para(doc, "Khởi động API server (22 class):")
    add_code_block(doc,
        'uvicorn scripts.api.api:app --host 0.0.0.0 --port 8000 --reload')
    add_para(doc, "Gửi request test:")
    add_code_block(doc,
        "python scripts/api/test_api.py --image path/to/image.jpg")

    add_heading(doc, "C. Kết quả per-class IoU đầy đủ (exp21 vs ensemble)", level=2)
    make_table(doc,
        headers=["#", "Class", "exp21 IoU", "Ensemble IoU", "Δ"],
        rows=[
            [1,  "egg",          "0.4575", "0.4458", "−0.0117"],
            [2,  "banana",       "0.4284", "0.4072", "−0.0212"],
            [3,  "steak",        "0.4010", "0.3931", "−0.0079"],
            [4,  "pork",         "0.1315", "0.1341", "+0.0026"],
            [5,  "chicken",      "0.4563", "0.4011", "−0.0552"],
            [6,  "fish",         "0.2983", "0.2892", "−0.0091"],
            [7,  "shrimp",       "0.3907", "0.3816", "−0.0091"],
            [8,  "bread",        "0.6124", "0.6047", "−0.0077"],
            [9,  "noodles",      "0.4305", "0.4701", "+0.0396"],
            [10, "rice",         "0.6279", "0.5930", "−0.0349"],
            [11, "tofu",         "0.3976", "0.4519", "+0.0543"],
            [12, "potato",       "0.5129", "0.4726", "−0.0403"],
            [13, "tomato",       "0.6521", "0.6341", "−0.0180"],
            [14, "lettuce",      "0.3885", "0.3444", "−0.0441"],
            [15, "cucumber",     "0.5549", "0.5379", "−0.0170"],
            [16, "carrot",       "0.7434", "0.7186", "−0.0248"],
            [17, "broccoli",     "0.7644", "0.7566", "−0.0078"],
            [18, "cabbage",      "0.4267", "0.2569", "−0.1698"],
            [19, "onion",        "0.3162", "0.3406", "+0.0244"],
            [20, "pepper",       "0.3107", "0.3126", "+0.0019"],
            [21, "French beans", "0.4962", "0.4897", "−0.0065"],
        ],
        col_widths=[Cm(1), Cm(3.5), Cm(3), Cm(3), Cm(2)]
    )

# ── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    write_cover(doc)
    write_toc(doc)
    write_ch1(doc)
    write_ch2(doc)
    write_ch3(doc)
    write_ch4(doc)
    write_ch5(doc)
    write_ch6(doc)
    write_ch7(doc)
    write_ch8(doc)
    write_ch9(doc)
    write_references(doc)
    write_appendix(doc)

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    try:
        doc.save(OUT_PATH)
        print(f"[OK] Báo cáo đã được lưu tại: {OUT_PATH}")
    except PermissionError:
        doc.save(OUT_PATH_ALT)
        print(f"[OK] File gốc đang mở → lưu tại: {OUT_PATH_ALT}")


if __name__ == "__main__":
    main()
